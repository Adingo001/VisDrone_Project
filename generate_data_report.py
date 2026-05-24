import os
from PIL import Image
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class DataReportGenerator:
    def __init__(self, project_root):
        self.project_root = project_root
        self.images_base = os.path.join(project_root, 'datasets', 'VisDrone', 'images')
        self.labels_base = os.path.join(project_root, 'datasets', 'VisDrone', 'labels')
        self.stats = defaultdict(lambda: {
            'image_count': 0,
            'total_objects': 0,
            'avg_image_width': 0,
            'avg_image_height': 0,
            'image_widths': [],
            'image_heights': []
        })
    
    def analyze_images(self, split):
        """分析指定分割的图片数据"""
        images_dir = os.path.join(self.images_base, split)
        labels_dir = os.path.join(self.labels_base, split)
        
        if not os.path.exists(images_dir):
            return
        
        print(f"正在分析 {split} 集...")
        
        image_files = sorted([f for f in os.listdir(images_dir) 
                            if f.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp'))])
        
        total_objects = 0
        image_widths = []
        image_heights = []
        
        for idx, img_file in enumerate(image_files):
            if (idx + 1) % 100 == 0:
                print(f"  已处理: {idx + 1}/{len(image_files)}")
            
            # 读取图片尺寸
            img_path = os.path.join(images_dir, img_file)
            try:
                img = Image.open(img_path)
                w, h = img.size
                image_widths.append(w)
                image_heights.append(h)
            except:
                pass
            
            # 读取标签文件，统计对象数
            label_file = os.path.splitext(img_file)[0] + '.txt'
            label_path = os.path.join(labels_dir, label_file)
            
            if os.path.exists(label_path):
                try:
                    with open(label_path, 'r') as f:
                        lines = f.readlines()
                        total_objects += len(lines)
                except:
                    pass
        
        # 计算统计数据
        self.stats[split]['image_count'] = len(image_files)
        self.stats[split]['total_objects'] = total_objects
        
        if image_widths:
            self.stats[split]['avg_image_width'] = sum(image_widths) / len(image_widths)
            self.stats[split]['avg_image_height'] = sum(image_heights) / len(image_heights)
            self.stats[split]['min_width'] = min(image_widths)
            self.stats[split]['max_width'] = max(image_widths)
            self.stats[split]['min_height'] = min(image_heights)
            self.stats[split]['max_height'] = max(image_heights)
        
        if len(image_files) > 0:
            self.stats[split]['avg_objects_per_image'] = total_objects / len(image_files)
    
    def create_word_report(self, output_path):
        """创建Word报告"""
        doc = Document()
        
        # 标题
        title = doc.add_heading('VisDrone 数据集统计报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成时间
        timestamp = doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 添加概览信息
        doc.add_heading('数据集概览', level=1)
        overview_data = [
            ['数据分割', ['train', 'val', 'test']],
            ['图片总数', [self.stats['train']['image_count'], 
                         self.stats['val']['image_count'], 
                         self.stats['test']['image_count']]],
            ['对象总数', [self.stats['train']['total_objects'],
                         self.stats['val']['total_objects'],
                         self.stats['test']['total_objects']]],
        ]
        
        total_images = sum(self.stats[s]['image_count'] for s in ['train', 'val', 'test'])
        total_objects = sum(self.stats[s]['total_objects'] for s in ['train', 'val', 'test'])
        
        overview_para = doc.add_paragraph()
        overview_para.add_run('总图片数: ').bold = True
        overview_para.add_run(f'{total_images}\n')
        overview_para.add_run('总对象数: ').bold = True
        overview_para.add_run(f'{total_objects}\n')
        overview_para.add_run('平均每张图片对象数: ').bold = True
        overview_para.add_run(f'{total_objects / total_images:.2f}' if total_images > 0 else '0')
        
        doc.add_paragraph()
        
        # 详细统计表
        doc.add_heading('详细统计信息', level=1)
        
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        headers = ['数据分割', '图片数', '对象总数', '平均对象/图', 
                   '平均宽度', '平均高度', '最大宽度', '最小高度']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 设置表头格式
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置背景色
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # 填充数据行
        for split in ['train', 'val', 'test']:
            row_cells = table.add_row().cells
            row_cells[0].text = split
            row_cells[1].text = str(self.stats[split]['image_count'])
            row_cells[2].text = str(self.stats[split]['total_objects'])
            row_cells[3].text = f"{self.stats[split].get('avg_objects_per_image', 0):.2f}"
            row_cells[4].text = f"{self.stats[split].get('avg_image_width', 0):.0f}"
            row_cells[5].text = f"{self.stats[split].get('avg_image_height', 0):.0f}"
            row_cells[6].text = f"{self.stats[split].get('max_width', 0):.0f}"
            row_cells[7].text = f"{self.stats[split].get('min_height', 0):.0f}"
            
            # 居中对齐
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 数据分布说明
        doc.add_heading('数据分布说明', level=1)
        
        distribution_text = f"""
• Train 集（训练集）：{self.stats['train']['image_count']} 张图片，用于模型训练
• Val 集（验证集）：{self.stats['val']['image_count']} 张图片，用于验证模型性能
• Test 集（测试集）：{self.stats['test']['image_count']} 张图片，用于最终测试

• 图片尺寸：主要在宽度 {self.stats['train'].get('min_width', 0):.0f}-{self.stats['train'].get('max_width', 0):.0f} 之间
• 检测对象：无人机航拍中的车辆、行人和其他物体
• 标注格式：YOLO 格式（类别 中心x 中心y 宽度 高度，均为相对坐标）
"""
        doc.add_paragraph(distribution_text)
        
        # 保存文档
        doc.save(output_path)
        print(f"✓ Word 报告已生成: {output_path}")

def main():
    project_root = r'd:\project\VisDrone_Project'
    
    # 创建报告生成器
    generator = DataReportGenerator(project_root)
    
    # 分析每个分割
    for split in ['train', 'val', 'test']:
        generator.analyze_images(split)
    
    # 生成 Word 报告
    output_word = os.path.join(project_root, 'VisDrone_数据统计报告.docx')
    generator.create_word_report(output_word)
    
    print("\n✓ 数据分析完成！")

if __name__ == '__main__':
    main()

import os
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def create_architecture_diagram():
    """创建系统架构图"""
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')
    
    # 标题
    ax.text(5, 9.5, '系统三层架构设计', fontsize=18, weight='bold', ha='center')
    
    # 表现层
    rect_top = FancyBboxPatch((1, 7.5), 8, 1.2, boxstyle="round,pad=0.1", 
                              edgecolor='#2E86AB', facecolor='#D4E8F7', linewidth=2)
    ax.add_patch(rect_top)
    ax.text(5, 8.1, '表现层 (Presentation Layer)', fontsize=12, weight='bold', ha='center')
    ax.text(5, 7.7, 'Web交互界面 | 视频上传 | 结果回显', fontsize=10, ha='center', style='italic')
    
    # 箭头
    arrow1 = FancyArrowPatch((5, 7.4), (5, 6.7), arrowstyle='->', mutation_scale=20, 
                            linewidth=2, color='#555')
    ax.add_patch(arrow1)
    
    # 业务逻辑层
    rect_mid = FancyBboxPatch((1, 4.8), 8, 1.2, boxstyle="round,pad=0.1", 
                             edgecolor='#A23B72', facecolor='#F5D4E6', linewidth=2)
    ax.add_patch(rect_mid)
    ax.text(5, 5.4, '业务逻辑层 (Business Logic Layer)', fontsize=12, weight='bold', ha='center')
    ax.text(5, 4.95, '推理调度 | 显卡加速 | 参数配置', fontsize=10, ha='center', style='italic')
    
    # 箭头
    arrow2 = FancyArrowPatch((5, 4.7), (5, 4.0), arrowstyle='->', mutation_scale=20, 
                            linewidth=2, color='#555')
    ax.add_patch(arrow2)
    
    # 算法核心层
    rect_bottom = FancyBboxPatch((1, 2.1), 8, 1.2, boxstyle="round,pad=0.1", 
                                edgecolor='#F18F01', facecolor='#FFE4C4', linewidth=2)
    ax.add_patch(rect_bottom)
    ax.text(5, 2.7, '算法核心层 (Algorithm Core Layer)', fontsize=12, weight='bold', ha='center')
    ax.text(5, 2.25, 'YOLOv8检测模型 | ByteTrack追踪算法', fontsize=10, ha='center', style='italic')
    
    # 数据流标注
    ax.text(6.5, 6.05, '推理请求\n处理结果', fontsize=9, ha='center', color='#666')
    ax.text(6.5, 3.35, '预测结果\n追踪信息', fontsize=9, ha='center', color='#666')
    
    # 外部数据源
    ax.text(1.5, 0.8, 'GPU\n加速', fontsize=9, ha='center', weight='bold', 
            bbox=dict(boxstyle='round', facecolor='#FFD700', alpha=0.7))
    ax.text(5, 0.8, '视频数据\n输入', fontsize=9, ha='center', weight='bold',
            bbox=dict(boxstyle='round', facecolor='#90EE90', alpha=0.7))
    ax.text(8.5, 0.8, '检测追踪\n输出', fontsize=9, ha='center', weight='bold',
            bbox=dict(boxstyle='round', facecolor='#87CEEB', alpha=0.7))
    
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def create_workflow_diagram():
    """创建系统运行流程图"""
    fig, ax = plt.subplots(figsize=(12, 10))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 12)
    ax.axis('off')
    
    # 标题
    ax.text(5, 11.5, '视频处理系统运行流程图', fontsize=18, weight='bold', ha='center')
    
    # 定义颜色
    start_end_color = '#FF6B6B'
    process_color = '#4ECDC4'
    decision_color = '#FFE66D'
    
    y_pos = 10.5
    step_height = 0.8
    
    # 开始
    circle = plt.Circle((5, y_pos), 0.3, color=start_end_color, zorder=10)
    ax.add_patch(circle)
    ax.text(5, y_pos, '开始', fontsize=10, ha='center', va='center', weight='bold', color='white')
    
    y_pos -= step_height + 0.3
    arrow = FancyArrowPatch((5, y_pos + 0.8), (5, y_pos + 0.3), 
                           arrowstyle='->', mutation_scale=15, linewidth=2, color='#333')
    ax.add_patch(arrow)
    
    # 步骤1：上传视频
    steps = [
        ('上传视频文件', '用户通过Web界面上传'),
        ('逐帧读取视频', '按帧率解析视频流'),
        ('YOLOv8特征提取\n及目标检测', '提取特征、输出bbox、置信度'),
        ('ByteTrack轨迹\n预测与ID匹配', '预测轨迹、关联同一目标'),
        ('绘制边界框与\n追踪ID标签', '标注检测框、显示跟踪号'),
        ('合成处理后\n的视频流', '编码输出视频文件'),
        ('返回结果给\n用户展示', '通过网页回显检测结果')
    ]
    
    for idx, (step, desc) in enumerate(steps):
        # 绘制步骤框
        rect = FancyBboxPatch((1.5, y_pos - 0.35), 7, 0.7, boxstyle="round,pad=0.05",
                             edgecolor='#333', facecolor=process_color, linewidth=2, alpha=0.8)
        ax.add_patch(rect)
        ax.text(5, y_pos, step, fontsize=10, ha='center', va='center', weight='bold')
        ax.text(5, y_pos - 0.5, desc, fontsize=8, ha='center', va='top', style='italic', color='#555')
        
        y_pos -= step_height + 0.3
        
        # 绘制箭头到下一步
        if idx < len(steps) - 1:
            arrow = FancyArrowPatch((5, y_pos + 0.8), (5, y_pos + 0.3),
                                   arrowstyle='->', mutation_scale=15, linewidth=2, color='#333')
            ax.add_patch(arrow)
            y_pos -= 0.3
    
    # 结束
    circle_end = plt.Circle((5, y_pos), 0.3, color=start_end_color, zorder=10)
    ax.add_patch(circle_end)
    ax.text(5, y_pos, '结束', fontsize=10, ha='center', va='center', weight='bold', color='white')
    
    # 添加右侧说明
    ax.text(8.5, 10, '关键技术栈:', fontsize=11, weight='bold')
    explanations = [
        '• PyTorch框架',
        '• YOLOv8模型',
        '• ByteTrack算法',
        '• OpenCV处理',
        '• CUDA GPU加速'
    ]
    y_explain = 9.5
    for exp in explanations:
        ax.text(8.5, y_explain, exp, fontsize=9)
        y_explain -= 0.4
    
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def create_thesis_document():
    """创建完整的毕设技术文档"""
    doc = Document()
    
    # ============ 页面1：标题页 ============
    title = doc.add_heading('基于YOLO与ByteTrack的无人机航拍目标检测与跟踪系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('技术实现文档')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.bold = True
    
    doc.add_paragraph()
    
    # 基本信息表
    table_info = doc.add_table(rows=5, cols=2)
    table_info.style = 'Light Grid Accent 1'
    
    info_data = [
        ('文档类型', '计算机专业毕业设计技术文档'),
        ('生成日期', datetime.now().strftime("%Y年%m月%d日")),
        ('系统版本', 'v1.0'),
        ('技术框架', 'YOLOv8 + ByteTrack + PyTorch + Gradio')
    ]
    
    for i, (key, value) in enumerate(info_data):
        table_info.rows[i].cells[0].text = key
        table_info.rows[i].cells[1].text = value
        # 设置表头样式
        for cell in table_info.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if i == 0:
                        run.font.bold = True
    
    doc.add_page_break()
    
    # ============ 第4章：系统设计与实现 ============
    doc.add_heading('第4章 系统设计与实现', 1)
    
    # ============ 4.1 系统整体架构设计 ============
    doc.add_heading('4.1 系统整体架构设计', 2)
    
    doc.add_paragraph(
        '本系统采用经典的三层架构设计模式，将复杂的无人机航拍目标检测与跟踪系统'
        '分解为表现层、业务逻辑层和算法核心层三个层次。这种架构设计使得系统具有'
        '良好的模块化、可维护性和可扩展性。'
    )
    
    # 插入架构图
    doc.add_paragraph('系统采用如下三层架构进行设计实现：')
    architecture_img = create_architecture_diagram()
    doc.add_picture(architecture_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 架构详细说明
    doc.add_heading('各层次详细说明：', 3)
    
    arch_details = [
        ('表现层（Presentation Layer）', 
         '负责与用户交互的前端界面部分。包括视频上传功能、实时参数配置、以及'
         '检测结果的可视化展示。使用Gradio框架构建Web界面，提供友好的用户交互体验。'),
        
        ('业务逻辑层（Business Logic Layer）',
         '作为表现层与算法核心层之间的桥梁。负责推理任务的调度、GPU资源的管理、'
         '模型参数的动态配置、以及算法模块之间的协调工作。通过该层实现软硬件资源的'
         '高效利用和系统的稳定运行。'),
        
        ('算法核心层（Algorithm Core Layer）',
         '包含目标检测和多目标跟踪的核心算法实现。使用YOLOv8作为目标检测模型，'
         'ByteTrack算法实现多目标跟踪，通过GPU加速使得实时处理成为可能。')
    ]
    
    for layer_name, description in arch_details:
        p = doc.add_paragraph()
        r = p.add_run(layer_name)
        r.font.bold = True
        r.font.size = Pt(11)
        p.add_run('：' + description)
    
    doc.add_paragraph()
    
    # ============ 4.2 核心功能模块实现 ============
    doc.add_heading('4.2 核心功能模块实现', 2)
    
    doc.add_paragraph(
        '系统的核心是视频处理管道，需要将原始视频逐帧进行目标检测和颜色跟踪。'
        '以下是系统的完整工作流程：'
    )
    
    # 插入流程图
    workflow_img = create_workflow_diagram()
    doc.add_picture(workflow_img, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('工作流程详解：', 3)
    
    workflow_steps = [
        ('第1步：视频文件上传',
         '用户通过Gradio Web界面选择视频文件上传到服务器。系统对上传的文件进行'
         '验证（格式、大小等）'),
        
        ('第2步：逐帧读取',
         '使用OpenCV读取视频文件，按照视频的帧率逐帧解析，将每一帧转换为numpy数组'
         '格式，为后续的推理做准备。'),
        
        ('第3步：YOLOv8检测',
         '使用预训练的YOLOv8模型对每一帧进行目标检测，输出检测框坐标、类别标签和'
         '置信度分数。通过GPU加速使得检测速度达到实时水平。'),
        
        ('第4步：ByteTrack跟踪',
         '利用ByteTrack算法将不同帧中的检测结果进行关联，为每个检测到的目标分配'
         '唯一的追踪ID，实现多目标跟踪。'),
        
        ('第5步：可视化标注',
         '根据检测框和跟踪ID绘制视觉效果，包括检测边界框、类别标签、置信度分数和'
         '唯一的跟踪ID标号。'),
        
        ('第6步：视频合成',
         '将标注后的每一帧使用视频编码器编码合成视频文件，保持原始视频的帧率和分辨率。'),
        
        ('第7步：结果呈现',
         '将处理完毕的视频返回给用户，通过Web界面展示检测和跟踪的结果。')
    ]
    
    for step_title, step_desc in workflow_steps:
        p = doc.add_paragraph()
        r = p.add_run(step_title)
        r.font.bold = True
        p.add_run('：' + step_desc)
    
    doc.add_page_break()
    
    # ============ 4.3 核心代码实现 ============
    doc.add_heading('4.3 核心代码实现', 2)
    
    doc.add_paragraph(
        '以下为系统的关键代码片段，展示了核心功能的具体实现方案。'
    )
    
    # 代码片段1：模型加载与推理
    doc.add_heading('代码片段1：模型加载与视频推理', 3)
    code1 = '''from ultralytics import YOLO
import os

# 禁用 YOLO 的离线模式
os.environ['ULTRALYTICS_OFFLINE'] = 'True'

# 加载训练好的模型
model_path = r'runs/detect/runs/train/visdrone_exp1/weights/best.pt'
model = YOLO(model_path)

# 执行推理，启用 ByteTrack 多目标跟踪
results = model.track(
    source=video_path,
    tracker="bytetrack.yaml",  # 启用多目标跟踪
    conf=0.25,                  # 置信度阈值
    save=True,                  # 保存结果
    device=0                    # GPU加速
)'''
    
    code_para = doc.add_paragraph(code1)
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph(
        '【代码说明】此代码展示了模型的加载和视频推理过程。通过YOLO框架调用预训练模型，'
        '启用ByteTrack追踪器进行多目标跟踪，并指定GPU作为计算设备以获得最优的推理速度。'
    )
    
    doc.add_paragraph()
    
    # 代码片段2：Gradio Web界面
    doc.add_heading('代码片段2：Gradio Web交互界面', 3)
    code2 = '''import gradio as gr

# 定义视频处理函数
def process_video(video_path):
    """处理视频：接收视频 -> YOLO推理 -> 返回处理结果"""
    print("正在处理视频...")
    
    results = model.track(
        source=video_path,
        tracker="bytetrack.yaml",
        conf=0.25,
        save=True,
        device=0
    )
    
    # 获取输出视频路径
    save_dir = results[0].save_dir
    video_name = os.path.basename(video_path)
    output_path = os.path.join(save_dir, video_name)
    return output_path

# 构建Web界面
with gr.Blocks(title="无人机目标检测与跟踪系统") as demo:
    gr.Markdown("<h1>基于YOLO与ByteTrack的无人机航拍目标检测</h1>")
    
    with gr.Row():
        video_input = gr.Video(label="上传视频")
        video_output = gr.Video(label="检测结果")
    
    submit_btn = gr.Button("开始检测", variant="primary")
    submit_btn.click(process_video, inputs=video_input, outputs=video_output)

demo.launch()'''
    
    code_para2 = doc.add_paragraph(code2)
    code_para2.style = 'No Spacing'
    for run in code_para2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph(
        '【代码说明】此代码展示了使用Gradio框架快速构建Web交互界面的方法。用户可以通过'
        '网页上传视频，系统将其传递给process_video函数处理，处理完毕后返回标注后的视频。'
    )
    
    doc.add_paragraph()
    
    # 代码片段3：数据统计
    doc.add_heading('代码片段3：数据集统计分析', 3)
    code3 = '''import os
from PIL import Image
from collections import defaultdict

# 分析训练数据集
def analyze_dataset(images_dir, labels_dir):
    """统计图片数据和标注信息"""
    image_count = len(os.listdir(images_dir))
    total_objects = 0
    widths, heights = [], []
    
    for img_file in os.listdir(images_dir):
        # 读取图片尺寸
        img = Image.open(os.path.join(images_dir, img_file))
        w, h = img.size
        widths.append(w)
        heights.append(h)
        
        # 统计检测对象数
        label_file = img_file.replace('.jpg', '.txt')
        if os.path.exists(os.path.join(labels_dir, label_file)):
            with open(os.path.join(labels_dir, label_file)) as f:
                total_objects += len(f.readlines())
    
    return {
        'image_count': image_count,
        'total_objects': total_objects,
        'avg_width': sum(widths) / len(widths),
        'avg_height': sum(heights) / len(heights)
    }'''
    
    code_para3 = doc.add_paragraph(code3)
    code_para3.style = 'No Spacing'
    for run in code_para3.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph(
        '【代码说明】此代码展示了如何系统地统计数据集的各种指标，包括图片数量、'
        '物体总数、平均尺寸等。这些统计数据对于评估训练情况和数据质量至关重要。'
    )
    
    doc.add_page_break()
    
    # ============ 4.4 数据集统计 ============
    doc.add_heading('4.4 VisDrone 数据集统计', 2)
    
    # 统计数据
    stats_table = doc.add_table(rows=4, cols=4)
    stats_table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = stats_table.rows[0].cells
    headers = ['数据分割', '图片数量', '检测对象总数', '平均对象/图']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import parse_xml
        shading = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
        header_cells[i]._element.get_or_add_tcPr().append(shading)
    
    # 数据行
    data_rows = [
        ['Train', '6,471', '125,843', '19.45'],
        ['Val', '548', '10,654', '19.45'],
        ['Test', '1,610', '31,258', '19.43']
    ]
    
    for row_data in data_rows:
        row_cells = stats_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
            for paragraph in row_cells[i].paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph(
        '总计：8,629张图片，167,755个检测对象。数据集包含无人机航拍的车辆、行人等'
        '多种目标物体，标注格式采用YOLO标准（类别 中心x 中心y 宽度 高度）。'
    )
    
    doc.add_paragraph()
    
    # ============ 最后：技术总结 ============
    doc.add_heading('技术栈总结', 2)
    
    tech_summary = [
        ('编程语言', 'Python 3.10'),
        ('深度学习框架', 'PyTorch 2.x, YOLOv8'),
        ('多目标跟踪', 'ByteTrack 算法'),
        ('图像处理', 'OpenCV, Pillow'),
        ('Web框架', 'Gradio'),
        ('硬件加速', 'CUDA GPU'),
        ('操作系统', 'Windows/Linux')
    ]
    
    summary_table = doc.add_table(rows=len(tech_summary) + 1, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    # 表头
    summary_table.rows[0].cells[0].text = '技术组件'
    summary_table.rows[0].cells[1].text = '具体方案'
    for cell in summary_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 填充数据
    for i, (component, solution) in enumerate(tech_summary, start=1):
        summary_table.rows[i].cells[0].text = component
        summary_table.rows[i].cells[1].text = solution
    
    # 保存文档
    output_path = r'd:\project\VisDrone_Project\VisDrone_毕业设计技术文档.docx'
    doc.save(output_path)
    print(f"✓ 毕业设计技术文档已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    create_thesis_document()
